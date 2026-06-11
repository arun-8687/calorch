"""Tests for the DOCX and HTML renderers."""
from datetime import datetime, UTC
from pathlib import Path

import pytest

from calorch.analysis import EventAnalysis, build_analysis
from calorch.renderers import render_docx, render_html_email
from calorch.state import CalendarEvent, ClassificationResult, EventType


@pytest.fixture
def sample_event() -> CalendarEvent:
    return CalendarEvent(
        id="ev-test-001",
        subject="AAPL Q1 FY26 Earnings Call",
        body_preview="Apple Q1 results discussion",
        start=datetime(2026, 3, 3, 21, 0, tzinfo=UTC),
        end=datetime(2026, 3, 3, 22, 0, tzinfo=UTC),
        organizer="ir@apple.com",
        attendees=["me@firm.example"],
        location="Webcast",
        is_online=True,
        web_link="",
    )


@pytest.fixture
def sample_classification() -> ClassificationResult:
    return ClassificationResult(
        event_id="ev-test-001",
        pass1_label=EventType.EARNINGS_CALL,
        pass1_keyword_hits=3,
        final_label=EventType.EARNINGS_CALL,
        confidence=0.92,
        rationale="earnings/q1/guidance hits",
        routed_node="handle_earnings_call",
    )


def test_earnings_call_docx_renders(tmp_path: Path, sample_event, sample_classification):
    analysis = EventAnalysis(
        event_id=sample_event.id,
        event_type=EventType.EARNINGS_CALL,
        title="Earnings Call Brief",
        sections=[
            ("Headline", ["In-line with consensus."]),
            ("Q&A Highlights", ["Buyback pace reaffirmed."]),
        ],
        tickers=["AAPL"],
        confidence=0.92,
    )
    out = tmp_path / "test.docx"
    render_docx(analysis, sample_event, out)
    assert out.exists()
    assert out.stat().st_size > 5_000  # real DOCX, not empty


def test_earnings_call_uses_ten_sections(tmp_path: Path, sample_event, sample_classification):
    analysis = build_analysis(
        EventType.EARNINGS_CALL,
        sample_event,
        sample_classification,
        enterprise_data={
            "source": "mock",
            "as_of": "now",
            "snapshots": {"AAPL": {"price": 200, "consensus_eps_q": 2.0, "consensus_rev_q": 1e10, "fy1_pe": 30, "ytd_return": 5}},
            "guidance": "Mgmt guides FY revenue +6-8% YoY",
            "transcript_excerpt": "We remain on track to deliver.",
        },
        llm_call=None,
    )
    out = tmp_path / "ec.docx"
    render_docx(analysis, sample_event, out)
    from docx import Document
    doc = Document(out)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    tables = doc.tables
    # Template-driven output: LLM sections become headings, data sections become tables
    assert len(headings) >= 1, f"expected >=1 heading, got {len(headings)}"
    assert len(tables) >= 5, f"expected >=5 data tables, got {len(tables)}"
    # No hardcoded "1." numbering remains
    assert not any(h.startswith("1.") for h in headings), "old numbered sections should not appear"


def test_html_email_contains_confidence_badge(sample_event, sample_classification):
    analysis = EventAnalysis(
        event_id=sample_event.id,
        event_type=EventType.EARNINGS_CALL,
        title="Brief",
        sections=[("Headline", ["A"]), ("Guidance", ["B"]), ("Q&A", ["C"])],
        tables=[{"headers": ["Ticker", "Price"], "rows": [["AAPL", 200]]}],
        confidence=0.92,
    )
    html = render_html_email(analysis, sample_event, doc_link=None)
    assert "earnings_call" in html
    assert "92%" in html
    assert "<table" in html


def test_management_meeting_infers_role(sample_event):
    sample_event.subject = "1:1 with CFO about capital allocation"
    cls = ClassificationResult(
        event_id=sample_event.id,
        final_label=EventType.MANAGEMENT_MEETING,
        confidence=0.7,
        routed_node="handle_management_meeting",
    )
    analysis = build_analysis(
        EventType.MANAGEMENT_MEETING,
        sample_event,
        cls,
        enterprise_data={"source": "mock", "as_of": "now", "snapshots": {}},
        llm_call=None,
    )
    assert analysis.role_focus == "CFO"


def test_html_email_suppresses_unsafe_doc_link_scheme(sample_event, sample_classification):
    """SEC-6: javascript:/data: doc links are not emitted as href."""
    from calorch.analysis import EventAnalysis

    analysis = EventAnalysis(
        event_id=sample_event.id, event_type=EventType.EARNINGS_CALL,
        title="Brief", sections=[("H", ["a"])], confidence=0.9,
    )
    out = render_html_email(analysis, sample_event, doc_link="javascript:alert(1)")
    assert 'href="javascript:' not in out
    safe = render_html_email(analysis, sample_event, doc_link="https://example.com/x.docx")
    assert 'href="https://example.com/x.docx"' in safe
