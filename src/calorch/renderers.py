"""Document and email generators.

These are the "rich output" half of the orchestrator:
  * python-docx generation with ten-section structure for earnings calls
    and event-type-specific layouts for the other seven.
  * HTML email builder with per-type templates, financial tables,
    confidence badges and inline CSS.

The LLM populates `EventAnalysis` and these renderers project that into
the final artefacts. The mock model produces deterministic text so the
demo runs without Azure OpenAI.
"""
from __future__ import annotations

import hashlib
import html
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from calorch.analysis import EventAnalysis  # re-exported for back-compat
from calorch.state import CalendarEvent, ClassificationResult, EventType
from calorch.telemetry import start_span

log = logging.getLogger("calorch.renderers")


# ---------------------------------------------------------------------------
# python-docx helpers
# ---------------------------------------------------------------------------
def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_thin_divider(doc: Document) -> None:
    """Add a thin horizontal line between sections (matching prep_scanner)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        if level == 1:
            run.font.size = Pt(14)


def _add_bullets(doc: Document, bullets: Iterable[str]) -> None:
    for b in bullets:
        if not b or not str(b).strip():
            continue
        doc.add_paragraph(str(b), style="List Bullet")


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = h
        for run in c.paragraphs[0].runs:
            run.bold = True
        _set_cell_shading(c, "D5E8F0")
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            table.cell(i, j).text = str(v)
    doc.add_paragraph("")


def _set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    style.paragraph_format.space_after = Pt(6)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Segoe UI"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    h1.paragraph_format.space_before = Pt(18)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Segoe UI"
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)


# ---------------------------------------------------------------------------
# Per-event-type renderers
# ---------------------------------------------------------------------------
def render_docx(analysis: EventAnalysis, event: CalendarEvent, out_path: Path) -> Path:
    with start_span("calorch.render.docx", event_id=event.id, event_type=analysis.event_type.value):
        return _render_docx_inner(analysis, event, out_path)


def _render_docx_inner(analysis: EventAnalysis, event: CalendarEvent, out_path: Path) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_doc_defaults(doc)

    # ---- title block (matching prep_scanner header style) ----
    title_text = analysis.title or "PREP PACK"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # Sub-header: event subject and date
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(event.subject)
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    date_str = event.start.strftime("%B %d, %Y | %I:%M %p IST") if hasattr(event.start, "strftime") else str(event.start)
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(date_str).italic = True

    _add_thin_divider(doc)

    if analysis.source_attribution:
        doc.add_paragraph().add_run(analysis.source_attribution).italic = True

    # ---- sections + tables (interleaved) ----
    ti = 0  # table index
    for heading, items in analysis.sections:
        _add_h(doc, heading, level=1)
        if items and items[0] == "__TABLE__":
            # Data section — render the next table here
            if ti < len(analysis.tables):
                t = analysis.tables[ti]
                title = t.get("title", "")
                if title:
                    tp = doc.add_paragraph()
                    tp.paragraph_format.space_before = Pt(8)
                    tr = tp.add_run(title)
                    tr.bold = True
                    tr.font.size = Pt(11)
                    tr.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
                _add_table(doc, t.get("headers", []), t.get("rows", []))
                ti += 1
        else:
            _add_bullets(doc, items)
    # Any remaining tables (from subsections, spare)
    while ti < len(analysis.tables):
        t = analysis.tables[ti]
        _add_table(doc, t.get("headers", []), t.get("rows", []))
        ti += 1

    # ---- data sources table ----
    if analysis.data_sources:
        _add_thin_divider(doc)
        _add_h(doc, "Data Sources", level=1)
        src_rows = [[s.get("source_name", ""), s.get("status", "").upper(), s.get("detail", "")]
                     for s in analysis.data_sources]
        _add_table(doc, ["Provider", "Status", "Detail"], src_rows)

    # ---- footer ----
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.paragraph_format.space_before = Pt(20)
    fr = foot.add_run(
        f"Generated {datetime.now(tz=timezone.utc).isoformat(timespec='seconds')}Z · "
        f"Confidence: {analysis.confidence:.0%} · Source: calorch (LangGraph)"
    )
    fr.italic = True
    fr.font.size = Pt(8)

    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# HTML email builder
# ---------------------------------------------------------------------------
_HTML_CSS = """
<style>
  body { font-family: -apple-system, 'Segoe UI', Arial, sans-serif; color: #1f2937; margin: 0; padding: 0; }
  .wrap { max-width: 640px; margin: 0 auto; }
  .header {
    background: linear-gradient(135deg, #1f3a5f 0%, #2e75b6 100%);
    color: white; padding: 18px 24px; border-radius: 8px 8px 0 0;
  }
  .header h1 { margin: 0; font-size: 20px; }
  .header .sub { font-size: 12px; opacity: 0.85; }
  .body { background: #ffffff; padding: 20px 24px; border: 1px solid #e5e7eb; border-top: 0; }
  .badge { display: inline-block; padding: 2px 10px; border-radius: 12px;
           font-size: 11px; font-weight: 600; background: #2e75b6; color: white; }
  .badge.high { background: #16a34a; }
  .badge.med  { background: #d97706; }
  .badge.low  { background: #dc2626; }
  .section { margin: 18px 0; }
  .section h2 { color: #1f3a5f; font-size: 15px; margin: 0 0 6px 0; }
  .section ul { margin: 6px 0 0 18px; padding: 0; }
  table.snap { border-collapse: collapse; width: 100%; font-size: 12px; }
  table.snap th, table.snap td { border: 1px solid #e5e7eb; padding: 6px 8px; text-align: right; }
  table.snap th { background: #f3f4f6; color: #1f3a5f; text-align: left; }
  table.snap td.tk { text-align: left; font-weight: 600; }
  .footer { font-size: 11px; color: #6b7280; text-align: center; padding: 12px; }
  a { color: #2e75b6; }
</style>
"""


def _confidence_badge(c: float) -> str:
    cls = "high" if c >= 0.75 else "med" if c >= 0.5 else "low"
    return f'<span class="badge {cls}">Confidence {c:.0%}</span>'


def render_html_email(analysis: EventAnalysis, event: CalendarEvent, doc_link: str | None, *, link_label: str = "Open DOCX") -> str:
    with start_span("calorch.render.html_email", event_id=event.id, event_type=analysis.event_type.value):
        return _render_html_email_inner(analysis, event, doc_link, link_label=link_label)


def _render_html_email_inner(analysis: EventAnalysis, event: CalendarEvent, doc_link: str | None, *, link_label: str = "Open DOCX") -> str:
    snap = analysis.tables[0] if analysis.tables else None
    snap_html = ""
    if snap:
        rows = "".join(
            "<tr>"
            + f'<td class="tk">{html.escape(str(r[0]))}</td>'
            + "".join(f"<td>{html.escape(str(c))}</td>" for c in r[1:])
            + "</tr>"
            for r in snap["rows"]
        )
        headers = "".join(f"<th>{html.escape(str(h))}</th>" for h in snap["headers"])
        snap_html = (
            '<table class="snap"><thead><tr>'
            + headers.replace("<th>", '<th colspan="1">', 0)
            + "</tr></thead><tbody>"
            + rows
            + "</tbody></table>"
        )

    sections_html = ""
    for heading, items in analysis.sections[:3]:
        bullets = "".join(f"<li>{html.escape(b)}</li>" for b in items)
        sections_html += (
            f'<div class="section"><h2>{html.escape(heading)}</h2><ul>{bullets}</ul></div>'
        )

    doc_link_html = (
        f'<p>Full brief: <a href="{html.escape(doc_link)}">{html.escape(link_label)}</a></p>'
        if doc_link
        else ""
    )

    return f"""<!doctype html>
<html><head><meta charset="utf-8">{_HTML_CSS}</head>
<body>
  <div class="wrap">
    <div class="header">
      <h1>{html.escape(analysis.title)}</h1>
      <div class="sub">{html.escape(event.subject)} · {event.start.strftime('%a %d %b %H:%M UTC')}</div>
    </div>
    <div class="body">
      <p>
        <span class="badge">{analysis.event_type.value}</span>
        {_confidence_badge(analysis.confidence)}
        {f'<span class="badge" style="background:#6b7280">{html.escape(analysis.role_focus)}</span>' if analysis.role_focus else ''}
      </p>
      {snap_html}
      {sections_html}
      {doc_link_html}
    </div>
    <div class="footer">Generated by calorch · LangGraph orchestrator · {datetime.now(tz=timezone.utc).isoformat(timespec='seconds')}Z</div>
  </div>
</body></html>"""


# ---------------------------------------------------------------------------
# Event-type analysis builders — return a populated EventAnalysis.
# Each builder asks the LLM for content, then we fall back to a typed
# template if the LLM is the mock model.
# ---------------------------------------------------------------------------
def build_analysis(
    event_type: EventType,
    event: CalendarEvent,
    cls: ClassificationResult,
    enterprise_data: dict[str, Any],
    llm_call,
    *,
    providers: Any = None,
    cik_lookup: Any = None,
) -> EventAnalysis:
    """Dispatch to the right builder by event type.

    ``providers`` is the calorch ``ProviderBundle``; if provided, the
    earnings/management/portfolio builders will pull real macro context
    (FRED/H.15), real segment splits (SEC iXBRL), and real guidance
    excerpts (SEC EFTS) for the first ticker on the event.
    """
    # Dispatch through the agent registry so each event type's builder is
    # declared in exactly one place (its calorch.agents module).
    from calorch.agents import get_agent

    with start_span("calorch.render.build_analysis", event_type=event_type.value) as span:
        return get_agent(event_type).analysis_builder(
            event, cls, enterprise_data, llm_call,
            providers=providers, cik_lookup=cik_lookup,
        )


# ---------------------------------------------------------------------------
# File helpers
# ---------------------------------------------------------------------------
def write_text(path: Path, content: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    return path


def sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()
